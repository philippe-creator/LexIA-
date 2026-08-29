"""
Rendu des documents juridiques (blocs neutres → fichier binaire).

Deux sorties depuis la même trame : DOCX (modifiable, brouillon relu par le
juriste) et PDF (version finale prête à imprimer/signer). La logique de contenu
vit dans generation/legal_documents.py ; ce module ne fait que la mise en forme.
"""

import io
import os

_FONTS_DIR = os.path.join(os.path.dirname(__file__), "fonts")


def render_docx(document: dict) -> bytes:
    """Rend un document (dict issu de build_document) en fichier DOCX."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    is_ar = document.get("lang") == "ar"

    def mark_rtl(paragraph):
        """Marque un paragraphe (et ses runs) comme texte bidirectionnel droite-
        à-gauche. Word gère nativement la forme des lettres arabes et l'ordre
        bidi une fois ce marqueur posé — pas besoin de retraiter le texte
        nous-mêmes, contrairement au PDF (voir render_pdf)."""
        pPr = paragraph._p.get_or_add_pPr()
        pPr.append(OxmlElement("w:bidi"))
        for run in paragraph.runs:
            rPr = run._element.get_or_add_rPr()
            rPr.append(OxmlElement("w:rtl"))
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            rFonts.set(qn("w:cs"), "Arial")

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for block in document["blocks"]:
        style, text = block["style"], block["text"]
        if style == "spacer":
            doc.add_paragraph("")
            continue
        p = doc.add_paragraph()
        run = p.add_run(text)
        if style == "title":
            run.bold = True
            run.font.size = Pt(15)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif style == "subtitle":
            run.bold = True
            run.font.size = Pt(12)
        elif style == "heading":
            run.bold = True
            run.font.size = Pt(11)
        elif style == "body_center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif style == "body_right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif style == "note":
            run.italic = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        else:  # body
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if is_ar:
            mark_rtl(p)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def render_pdf(document: dict) -> bytes:
    """Rend un document (dict issu de build_document) en fichier PDF.

    FR/EN : polices core de fpdf2 (Helvetica), encodées en latin-1 — couvrent
    les accents français/anglais.

    AR : Helvetica ne contient aucun glyphe arabe et fpdf2 ne fait pas la mise
    en forme contextuelle des lettres arabes ni la réorganisation bidirection-
    nelle tout seul. On embarque donc une police Unicode arabe (Amiri, licence
    SIL OFL — voir generation/fonts/OFL.txt), et pour chaque bloc de texte :
    1) `arabic_reshaper` choisit la forme contextuelle de chaque lettre arabe
       (une lettre arabe change de forme selon sa position dans le mot — un
       rendu naïf glyphe-par-glyphe produit des lettres isolées illisibles) ;
    2) `python-bidi` réordonne le texte en ordre visuel gauche-à-droite, seul
       ordre que fpdf2 sait afficher tel quel.
    Le texte résultant est ensuite aligné à droite, la lecture visuelle
    correspondant bien à un rendu arabe lu de droite à gauche.
    """
    from fpdf import FPDF

    is_ar = document.get("lang") == "ar"

    def latin1(s: str) -> str:
        return (s.replace("’", "'").replace("‘", "'")
                 .replace("“", '"').replace("”", '"')
                 .replace("–", "-").replace("—", "-")
                 .replace("…", "...")
                 .encode("latin-1", "replace").decode("latin-1"))

    def wrap_ar_lines(pdf, text: str, width: float) -> list:
        """Découpe le texte arabe LOGIQUE en lignes qui tiennent dans `width`,
        puis ne met en forme (reshape) et ne réordonne (bidi) que chaque ligne
        individuellement.

        Appliquer le reshape + bidi sur le paragraphe ENTIER avant le retour à
        la ligne automatique de fpdf2 casse l'ordre des lignes : bidi.get_display()
        produit une chaîne en ordre visuel gauche-à-droite unique, et le
        word-wrap de fpdf2 découpe alors cette chaîne comme si elle était de
        l'anglais normal — la ligne qui doit apparaître EN BAS du paragraphe se
        retrouve rendue EN HAUT. En calculant les coupures de ligne sur le texte
        logique (police déjà réglée, mesure via reshape seul, sans bidi), l'ordre
        des lignes reste correct ; bidi n'intervient qu'ensuite, ligne par ligne,
        pour l'affichage interne de chacune."""
        import arabic_reshaper
        from bidi.algorithm import get_display
        words = text.split(" ")
        lines, current = [], ""
        for w in words:
            candidate = f"{current} {w}".strip() if current else w
            if not current or pdf.get_string_width(arabic_reshaper.reshape(candidate)) <= width:
                current = candidate
            else:
                lines.append(current)
                current = w
        if current:
            lines.append(current)
        return [get_display(arabic_reshaper.reshape(line)) for line in lines]

    pdf = FPDF(format="A4")
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.set_margins(left=20, top=20, right=20)
    pdf.add_page()
    width = pdf.w - pdf.l_margin - pdf.r_margin

    font = "Helvetica"
    if is_ar:
        font = "Amiri"
        pdf.add_font(font, "", os.path.join(_FONTS_DIR, "Amiri-Regular.ttf"))
        pdf.add_font(font, "B", os.path.join(_FONTS_DIR, "Amiri-Bold.ttf"))
        # Pas de variante italique embarquée pour Amiri — le style "note"
        # retombe donc sur le regular en arabe (voir plus bas).

    def draw(style_flag, size, align, line_height, raw_text):
        """Affiche un bloc de texte, ligne par ligne pour l'arabe (voir
        wrap_ar_lines), en un seul multi_cell sinon."""
        pdf.set_font(font, style_flag, size)
        if is_ar:
            for line in wrap_ar_lines(pdf, raw_text, width):
                pdf.set_xy(pdf.l_margin, pdf.get_y())
                pdf.multi_cell(width, line_height, line, align=align)
        else:
            pdf.set_xy(pdf.l_margin, pdf.get_y())
            pdf.multi_cell(width, line_height, latin1(raw_text), align=align)

    for block in document["blocks"]:
        style, raw_text = block["style"], block["text"]
        if style == "spacer":
            pdf.ln(4)
            continue
        # fpdf2 (2.8.7) a un bug reproductible : quand un multi_cell() qui vient
        # de replier son texte sur 2 lignes est immédiatement suivi d'un autre
        # multi_cell(), le second voit son texte tronqué dans le flux PDF rendu
        # (le texte Python, lui, est intact — le bug est interne à fpdf2). Forcer
        # la position avant CHAQUE multi_cell neutralise le problème. Reproduit
        # et corrigé sur « Il/Elle est en poste depuis le... » (attestation de
        # travail) qui sortait tronqué à « Il/Elle est e ».
        pdf.set_xy(pdf.l_margin, pdf.get_y())
        default_align = "R" if is_ar else "L"
        if style == "title":
            draw("B", 15, "C", 8, raw_text)
        elif style == "subtitle":
            draw("B", 12, default_align, 6, raw_text)
        elif style == "heading":
            draw("B", 11, default_align, 6, raw_text)
        elif style == "body_center":
            draw("", 11, "C", 6, raw_text)
        elif style == "body_right":
            draw("", 11, "R", 6, raw_text)
        elif style == "note":
            pdf.ln(2)
            pdf.set_text_color(128, 128, 128)
            draw("" if is_ar else "I", 8, default_align, 4, raw_text)
            pdf.set_text_color(0, 0, 0)
        else:  # body
            # align="J" (justifié) a un bug de rendu dans fpdf2 sur certaines
            # phrases courtes tenant sur une seule ligne : le texte est purement
            # perdu (pas juste mal réparti) — reproduit avec « Il/Elle est en
            # poste depuis le... » dans l'attestation de travail. Alignement à
            # gauche (ou à droite en arabe) : moins "propre" visuellement, mais
            # fiable.
            draw("", 11, default_align, 6, raw_text)

    out = pdf.output()  # fpdf2 ≥ 2.8 renvoie un bytearray
    return bytes(out)
