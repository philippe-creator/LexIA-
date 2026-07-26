import io
import json
from pathlib import Path
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from core.database import get_db
from api.core.dependencies import CurrentUser, require_feature
from core.plans import FEATURE_EXPORT
from api.repositories.conversation_repo import ConversationRepository

router = APIRouter(prefix="/export", tags=["Export"])

@router.get("/conversations/{conv_id}/json")
async def export_json(conv_id: str, current_user: CurrentUser, db: Session = Depends(get_db),
                      _gate=Depends(require_feature(FEATURE_EXPORT))):
    conv = ConversationRepository(db).get(conv_id, current_user.id)
    if not conv: raise HTTPException(404, "Conversation introuvable.")
    messages, _ = ConversationRepository(db).get_history(conv.id, limit=1000)
    data = {
        "conversation": {"id": conv.id, "title": conv.title, "domain": conv.domain, "created_at": conv.created_at.isoformat(), "updated_at": conv.updated_at.isoformat()},
        "messages": [{"role": m.role, "content": m.content, "citations": m.citations or [], "confidence_score": m.confidence_score, "created_at": m.created_at.isoformat()} for m in messages],
    }
    buf = io.BytesIO(json.dumps(data, ensure_ascii=False, indent=2).encode("utf-8"))
    return StreamingResponse(buf, media_type="application/json", headers={"Content-Disposition": f"attachment; filename=conversation_{conv.id}.json"})

@router.get("/conversations/{conv_id}/docx")
async def export_docx(conv_id: str, current_user: CurrentUser, db: Session = Depends(get_db),
                      _gate=Depends(require_feature(FEATURE_EXPORT))):
    conv = ConversationRepository(db).get(conv_id, current_user.id)
    if not conv: raise HTTPException(404, "Conversation introuvable.")
    messages, _ = ConversationRepository(db).get_history(conv.id, limit=1000)
    from docx import Document
    doc = Document()
    doc.add_heading(f"Conversation — {conv.title or 'Sans titre'}", level=1)
    doc.add_paragraph(f"Domaine : {conv.domain or 'N/A'} | Créée le : {conv.created_at.strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph("")
    for m in messages:
        p = doc.add_paragraph()
        p.add_run(f"{'Utilisateur' if m.role == 'user' else 'Assistant'}: ").bold = True
        p.add_run(m.content)
        if m.citations:
            doc.add_paragraph("Sources : " + ", ".join(c.get("label", "") for c in (m.citations or [])[:3])).italic = True
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f"attachment; filename=conversation_{conv.id}.docx"})
