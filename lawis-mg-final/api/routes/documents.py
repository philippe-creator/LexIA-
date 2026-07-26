import hashlib, uuid
from pathlib import Path
from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from sqlalchemy.orm import Session
from core.config import settings
from core.database import get_db, UserDocument
from api.core.dependencies import CurrentUser, require_feature
from core.domains import DOMAINS
from core.plans import FEATURE_CONTRACT_AUDIT

router = APIRouter(prefix="/documents", tags=["Documents"])
UPLOADS_DIR = Path(settings.UPLOADS_DIR)

# Matières qu'un audit de contrat de travail doit couvrir. Chacune fait l'objet
# d'une recherche distincte : une requête unique et large ramenait des passages
# tous centrés sur le même thème et manquait des articles décisifs.
AUDIT_TOPICS = [
    "durée maximale de la période d'essai et son renouvellement",
    "délai de préavis en cas de rupture du contrat de travail",
    "salaire minimum légal et modalités de paiement de la rémunération",
    "congés payés annuels du salarié",
    "indemnité de licenciement et rupture abusive du contrat",
    "durée légale du travail hebdomadaire et heures supplémentaires",
]
ALLOWED_MIME = {"application/pdf","application/msword","application/vnd.openxmlformats-officedocument.wordprocessingml.document","text/plain","text/html"}
MIME_EXTENSIONS = {
    "application/pdf": {".pdf"},
    "application/msword": {".doc"},
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": {".docx"},
    "text/plain": {".txt"},
    "text/html": {".html", ".htm"},
}
MAGIC_SIGNATURES = {
    "application/pdf": (b"%PDF-",),
    "application/msword": (b"\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1",),
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": (b"PK\x03\x04",),
}

def _content_matches(data: bytes, content_type: str) -> bool:
    """Vérifie les premiers octets réels du fichier plutôt que de faire confiance au Content-Type déclaré par le client."""
    sigs = MAGIC_SIGNATURES.get(content_type)
    if sigs: return any(data.startswith(sig) for sig in sigs)
    if content_type == "text/plain": return b"\x00" not in data[:2048]
    if content_type == "text/html":
        head = data[:512].lstrip().lower()
        return head.startswith(b"<!doctype") or b"<html" in head
    return False

@router.post("/upload", status_code=201)
async def upload(file: UploadFile = File(...), domain: str = Form("divers"), current_user: CurrentUser = None, db: Session = Depends(get_db)):
    if domain not in DOMAINS: raise HTTPException(400, f"Domaine invalide : {domain}")
    if file.content_type not in ALLOWED_MIME: raise HTTPException(400, f"Type non supporté : {file.content_type}")
    ext = Path(file.filename).suffix.lower()
    allowed_exts = {f".{e.lstrip('.')}" for e in settings.ALLOWED_EXTENSIONS}
    if ext not in allowed_exts:
        raise HTTPException(400, f"Extension non autorisée : {ext}")
    if ext not in MIME_EXTENSIONS.get(file.content_type, set()):
        raise HTTPException(400, "L'extension du fichier ne correspond pas au type déclaré.")
    data = await file.read()
    if len(data) > settings.max_upload_bytes: raise HTTPException(413, f"Fichier trop volumineux (max {settings.MAX_UPLOAD_SIZE_MB}MB)")
    if not _content_matches(data, file.content_type):
        raise HTTPException(400, "Le contenu du fichier ne correspond pas au type déclaré.")
    checksum = hashlib.sha256(data).hexdigest()
    existing = db.query(UserDocument).filter(UserDocument.user_id==current_user.id, UserDocument.checksum==checksum).first()
    if existing: raise HTTPException(409, f"Document déjà importé : {existing.original_filename}")
    safe_name = f"{uuid.uuid4()}{ext}"
    user_dir = UPLOADS_DIR / current_user.id
    user_dir.mkdir(parents=True, exist_ok=True)
    (user_dir / safe_name).write_bytes(data)
    doc = UserDocument(user_id=current_user.id, filename=safe_name, original_filename=file.filename, file_path=str(user_dir/safe_name), file_size=len(data), mime_type=file.content_type, domain=domain, checksum=checksum, status="pending")
    db.add(doc); db.commit(); db.refresh(doc)
    try: await _index(doc, db)
    except Exception as e: doc.status="error"; doc.error_message=str(e); db.commit()
    db.refresh(doc)
    return {"id":doc.id,"filename":doc.original_filename,"status":doc.status,"domain":doc.domain,"chunk_count":doc.chunk_count,"file_size":doc.file_size}

async def _index(doc, db):
    from ingestion.ocr.pdf_extractor import extract_text
    from ingestion.pipeline import ingest_text
    from datetime import datetime
    doc.status="processing"; db.commit()
    p = Path(doc.file_path)
    page_offsets = None
    if doc.mime_type=="application/pdf":
        extracted = extract_text(p)
        text = extracted["text"]
        page_offsets = extracted.get("page_offsets")
    elif doc.mime_type=="text/plain": text = p.read_text(encoding="utf-8",errors="ignore")
    else:
        try:
            import docx; text="\n".join(para.text for para in docx.Document(str(p)).paragraphs)
        except: text=""
    if not text or len(text.strip())<50: raise ValueError("Texte extrait vide.")
    # source="user_upload" : le pipeline exclut délibérément ces documents du
    # snapshot de version partagé (voir ingestion/pipeline.py) pour ne pas
    # exposer un upload privé via /compare à d'autres utilisateurs.
    result = ingest_text(text, domain=doc.domain, filename=doc.original_filename, source="user_upload", extra_metadata={"user_id":doc.user_id,"document_id":doc.id}, page_offsets=page_offsets)
    doc.status="indexed"; doc.chunk_count=result["chunks_indexed"]; doc.indexed_at=datetime.utcnow(); db.commit()

@router.get("/")
async def list_docs(current_user: CurrentUser, db: Session = Depends(get_db)):
    docs = db.query(UserDocument).filter(UserDocument.user_id==current_user.id).order_by(UserDocument.created_at.desc()).all()
    return [{"id":d.id,"filename":d.original_filename,"domain":d.domain,"status":d.status,"chunk_count":d.chunk_count,"file_size":d.file_size,"error_message":d.error_message,"created_at":d.created_at.isoformat()} for d in docs]

@router.post("/{doc_id}/audit")
async def audit_document(doc_id: str, current_user: CurrentUser, db: Session = Depends(get_db),
                         _gate=Depends(require_feature(FEATURE_CONTRACT_AUDIT))):
    """Audite un contrat importé au regard du droit du travail marocain :
    extrait son texte, récupère des passages de loi pertinents (grounding RAG)
    et génère un rapport structuré citant les articles applicables."""
    import asyncio
    from ingestion.ocr.pdf_extractor import extract_text
    from retrieval.hybrid_retriever import retrieve
    from generation.prompt_builder import build_audit_prompt, format_citations
    from generation.llm_client import generate

    doc = db.query(UserDocument).filter(UserDocument.id == doc_id, UserDocument.user_id == current_user.id).first()
    if not doc:
        raise HTTPException(404, "Document introuvable.")

    p = Path(doc.file_path)
    if not p.exists():
        raise HTTPException(404, "Fichier introuvable sur le serveur.")
    if doc.mime_type == "application/pdf":
        text = extract_text(p)["text"]
    elif doc.mime_type == "text/plain":
        text = p.read_text(encoding="utf-8", errors="ignore")
    else:
        try:
            import docx; text = "\n".join(par.text for par in docx.Document(str(p)).paragraphs)
        except Exception:
            text = ""
    if not text or len(text.strip()) < 50:
        raise HTTPException(422, "Impossible d'extraire un texte exploitable de ce document.")

    loop = asyncio.get_event_loop()
    # Grounding : on ancre l'audit sur des articles réels du corpus travail.
    # user_id=None → uniquement les textes officiels partagés, jamais les uploads
    # privés (le contrat audité ne doit pas se citer lui-même comme référence).
    #
    # Une requête unique et large ne convenait pas : les 6 passages retenus se
    # concentraient sur une même matière et laissaient de côté des articles
    # décisifs. L'article 14 (durée maximale de la période d'essai) n'était
    # jamais récupéré, et l'audit concluait à tort qu'une période d'essai de
    # 6 mois est conforme « faute de durée maximale dans le code ». On interroge
    # donc chaque matière qu'un audit de contrat doit couvrir.
    chunks, seen = [], set()
    for topic in AUDIT_TOPICS:
        topic_chunks, _, _, _ = await loop.run_in_executor(
            None, lambda q=topic: retrieve(query=q, top_k=3, forced_domains=["travail"], user_id=None)
        )
        for c in topic_chunks:
            key = (c.get("text") or "")[:120]
            if key and key not in seen:
                seen.add(key)
                chunks.append(c)
    # Le budget de contexte (tiktoken) élague ensuite si nécessaire.
    system_prompt, user_message = build_audit_prompt(text, chunks)
    try:
        report = await loop.run_in_executor(None, lambda: generate(system_prompt, user_message))
    except Exception as e:
        from loguru import logger
        logger.error(f"Erreur audit contrat : {e}")
        raise HTTPException(503, "Le service d'analyse est momentanément indisponible. Réessayez.")
    return {"document_id": doc.id, "filename": doc.original_filename, "report": report, "citations": format_citations(chunks)}

@router.delete("/{doc_id}")
async def delete_doc(doc_id: str, current_user: CurrentUser, db: Session = Depends(get_db)):
    doc = db.query(UserDocument).filter(UserDocument.id==doc_id, UserDocument.user_id==current_user.id).first()
    if not doc: raise HTTPException(404,"Document introuvable.")
    try: Path(doc.file_path).unlink(missing_ok=True)
    except: pass
    db.delete(doc); db.commit()
    return {"message":"Supprimé."}
